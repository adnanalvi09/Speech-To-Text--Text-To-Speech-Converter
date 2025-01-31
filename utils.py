import docx
import PyPDF2
import speech_recognition as sr
from docx import Document
import pyttsx3
import streamlit as st



def text_to_speech(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()




def recognize_speech(placeholder):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        placeholder.write("Speak something...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)
    try:
        placeholder.write("Recognizing...")
        text = recognizer.recognize_google(audio)
        return text
    except sr.UnknownValueError:
        placeholder.write("Sorry, could not understand audio.")
        return ""
    except sr.RequestError as e:
        placeholder.write("Could not request results from Google Speech Recognition service; {0}".format(e))
        return ""

def write_to_docx(text, placeholder ,filename="output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    placeholder.write(f"Text has been written to {filename}")



def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfFileReader(file)
        for page_num in range(reader.numPages):
            page = reader.getPage(page_num)
            text += page.extractText()
    return text




